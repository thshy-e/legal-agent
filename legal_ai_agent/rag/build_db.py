# knowledge_base_builder.py
"""
劳动法智能知识库构建脚本
功能：
1. 从文本列表或外部文件构建向量库（法条库 / 案例库）
2. 持久化到本地磁盘（ChromaDB）
3. 支持动态增量添加文档
4. 支持按 ID 删除文档
5. 提供检索接口（供 Agent 调用）

依赖：
pip install langchain-chroma langchain-community chromadb langchain-openai sentence-transformers
"""

import os
import re
from typing import List, Optional, Union
from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import DashScopeEmbeddings
from legal_ai_agent.config.settings import (
    CASE_DATA_DIR,
    CHROMA_DIR,
    EMBEDDING_MODEL,
    LAW_DATA_DIR,
    RAW_DATA_DIR,
    get_dashscope_api_key,
)

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

# ========== 配置参数 ==========
PERSIST_DIR = str(CHROMA_DIR)  # 持久化目录

# 示例：本地 Embedding 模型（无需联网，比赛推荐）也可替换为 DashScope
try:
    from langchain_huggingface import HuggingFaceEmbeddings
    EMBEDDING_MODEL_NAME = "BAAI/bge-small-zh-v1.5"   # 轻量且效果好的中文向量模型
    embedding_function = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL_NAME)
except ImportError:
    # 备用：阿里云 DashScope 向量服务
    embedding_function = DashScopeEmbeddings(
        dashscope_api_key=get_dashscope_api_key(),
        model=EMBEDDING_MODEL or "text-embedding-v3"
    )

# ========== 文本切分策略 ==========
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=3000,
    chunk_overlap=200,
    separators=["\n\n", "\n", "。", "；", " ", ""],
)

ARTICLE_PATTERN = re.compile(r"(第[一二三四五六七八九十百零〇\d]+条)")


def _chinese_number_to_int(value: str) -> str:
    if not value:
        return ""
    if value.isdigit():
        return value

    digits = {"零": 0, "〇": 0, "一": 1, "二": 2, "三": 3, "四": 4, "五": 5, "六": 6, "七": 7, "八": 8, "九": 9}
    units = {"十": 10, "百": 100, "千": 1000}
    total = 0
    current = 0
    for char in value:
        if char in digits:
            current = digits[char]
        elif char in units:
            unit = units[char]
            if current == 0:
                current = 1
            total += current * unit
            current = 0
    total += current
    return str(total) if total else value


def _article_metadata(text: str, law_name: str) -> dict:
    metadata = {"law_name": law_name}
    match = ARTICLE_PATTERN.search(text or "")
    if match:
        metadata["article_label"] = match.group(1)
        raw_article = match.group(1).replace("第", "").replace("条", "")
        metadata["article"] = _chinese_number_to_int(raw_article)
    return metadata

# ========== 核心函数：构建/加载向量库 ==========
class _BatchEmbeddings:
    """包装 Embeddings，限制每次 embed_documents 的批量大小，避免 DashScope API 超时"""
    def __init__(self, embeddings, batch_size: int = 5):
        self._embeddings = embeddings
        self._batch_size = batch_size

    def embed_documents(self, texts):
        results = []
        for i in range(0, len(texts), self._batch_size):
            batch = texts[i:i+self._batch_size]
            results.extend(self._embeddings.embed_documents(batch))
        return results

    def embed_query(self, text):
        return self._embeddings.embed_query(text)

embedding_function = _BatchEmbeddings(embedding_function, batch_size=20)

def build_law_kb(law_texts: Union[str, List[str]], collection_name: str = "labor_law") -> Chroma:
    """构建法条向量库（覆盖原有数据）"""
    docs = _texts_to_docs(law_texts, law_name="劳动法")
    db = Chroma.from_documents(
        docs,
        embedding_function,
        collection_name=collection_name,
        persist_directory=PERSIST_DIR
    )
    return db

def build_case_kb(case_texts: Union[str, List[str]], collection_name: str = "labor_cases") -> Chroma:
    """构建案例向量库（覆盖原有数据）"""
    docs = _texts_to_docs(case_texts)
    db = Chroma.from_documents(
        docs,
        embedding_function,
        collection_name=collection_name,
        persist_directory=PERSIST_DIR
    )
    return db


def load_kb(collection_name: str) -> Optional[Chroma]:
    """加载已持久化的向量库，不存在则返回 None"""
    if not os.path.exists(PERSIST_DIR):
        return None
    try:
        db = Chroma(
            collection_name=collection_name,
            embedding_function=embedding_function,
            persist_directory=PERSIST_DIR
        )
        return db
    except Exception:
        return None

def add_docs_to_kb(texts: Union[str, List[str]], collection_name: str, metadatas: Optional[List[dict]] = None):
    """增量添加文档到现有向量库"""
    db = load_kb(collection_name)
    if db is None:
        db = Chroma(
            collection_name=collection_name,
            embedding_function=embedding_function,
            persist_directory=PERSIST_DIR
        )
    docs = _texts_to_docs(texts) if metadatas is None else _texts_to_docs(texts, metadatas)
    db.add_documents(docs)

def delete_docs_from_kb(ids: List[str], collection_name: str):
    """按文档 ID 删除文档"""
    db = load_kb(collection_name)
    if db:
        db.delete(ids=ids)

def retrieve_from_kb(query: str, collection_name: str, k: int = 4) -> str:
    """从指定知识库检索相关文档，返回拼接字符串"""
    db = load_kb(collection_name)
    if db is None:
        return "暂无相关知识库数据。"
    docs = db.similarity_search(query, k=k)
    if not docs:
        return "未检索到相关内容。"
    return "\n\n".join([f"【参考 {i+1}】{doc.page_content}" for i, doc in enumerate(docs)])

# ========== 内部工具函数 ==========
def _texts_to_docs(
    texts: Union[str, List[str]],
    metadatas: Optional[List[dict]] = None,
    law_name: Optional[str] = None,
) -> List[Document]:
    """将原始文本切分为 LangChain Document 列表"""
    if isinstance(texts, str):
        texts = [texts]
    all_docs = []
    for idx, text in enumerate(texts):
        chunks = text_splitter.split_text(text)
        for chunk in chunks:
            if metadatas and idx < len(metadatas):
                all_docs.append(Document(page_content=chunk, metadata=metadatas[idx]))
            elif law_name:
                all_docs.append(Document(page_content=chunk, metadata=_article_metadata(chunk, law_name)))
            else:
                all_docs.append(Document(page_content=chunk))
    return all_docs

# ========== Docx 文件读取工具 ==========
DOCX_DIR = str(RAW_DATA_DIR)

def read_docx_text(file_path: str) -> str:
    """读取 .docx 文件并返回全部文本内容"""
    if DocxDocument is None:
        raise ImportError("python-docx 未安装，请运行: pip install python-docx")
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")
    doc = DocxDocument(file_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)

def read_docx_split_by_articles(file_path: str) -> List[str]:
    """读取 .docx 文件并按法条编号（第X条）拆分为段落列表"""
    full_text = read_docx_text(file_path)
    articles = re.split(r"(?=第[一二三四五六七八九十百零〇\d]+条)", full_text)
    return [a.strip() for a in articles if a.strip()]

def read_docx_split_by_cases(file_path: str) -> List[str]:
    """读取 .docx 文件并按案例标题拆分为案例列表"""
    full_text = read_docx_text(file_path)
    parts = full_text.split("\n【基本案情】")
    if len(parts) <= 1:
        return [full_text]
    cases = []
    for i, part in enumerate(parts):
        if i == 0:
            lines = part.strip().split("\n")
            title_lines = [l.strip() for l in lines if l.strip()]
            title = " ".join(title_lines)
        else:
            prev = cases[-1] if cases else ""
            if prev:
                rest_lines = part.strip().split("\n")
                title = "【基本案情】" + rest_lines[0] if rest_lines else ""
                body = "\n".join(rest_lines[1:]) if len(rest_lines) > 1 else ""
                full = prev + "\n" + title
                if body:
                    full += "\n" + body
            else:
                full = "【基本案情】" + part.strip()
            cases[-1] = full
            continue
        rest_lines = part.strip().split("\n")
        body = "\n".join(rest_lines[len(title_lines):])
        full = title
        if body:
            full += "\n" + body
        cases.append(full)
    return [c for c in cases if len(c) > 100]

# ========== 从 Docx 文件加载知识库 ==========
LAW_DOCX_PATH = os.path.join(DOCX_DIR, "劳动法.docx")
CASE_DOCX_PATH = os.path.join(DOCX_DIR, "劳动法经典案例.docx")

def load_law_from_docx() -> List[str]:
    """从劳动法.docx 读取法条内容"""
    if not os.path.exists(LAW_DOCX_PATH):
        print(f"⚠️ 未找到 {LAW_DOCX_PATH}，尝试读取 data/law/ 下的文本文件...")
        law_dir = str(LAW_DATA_DIR)
        texts = []
        if os.path.exists(law_dir):
            for fname in os.listdir(law_dir):
                if fname.endswith(".txt"):
                    with open(os.path.join(law_dir, fname), "r", encoding="utf-8") as f:
                        texts.append(f.read())
        return texts if texts else []
    articles = read_docx_split_by_articles(LAW_DOCX_PATH)
    print(f"  从劳动法.docx 读取到 {len(articles)} 条法条")
    return articles

def load_cases_from_docx() -> List[str]:
    """从劳动法经典案例.docx 读取案例内容"""
    if not os.path.exists(CASE_DOCX_PATH):
        print(f"  未找到 {CASE_DOCX_PATH}，尝试读取 data/case/ 下的文本文件...")
        case_dir = str(CASE_DATA_DIR)
        texts = []
        if os.path.exists(case_dir):
            for fname in os.listdir(case_dir):
                if fname.endswith(".txt"):
                    with open(os.path.join(case_dir, fname), "r", encoding="utf-8") as f:
                        texts.append(f.read())
        return texts if texts else []
    full_text = read_docx_text(CASE_DOCX_PATH)
    total_chars = len(full_text)
    print(f"  从劳动法经典案例.docx 读取到 {total_chars} 字符")
    return [full_text]

# ========== 一键初始化 ==========
def init_all_kb(force_rebuild: bool = False):
    """初始化法条库和案例库（若已存在则不重建，除非 force_rebuild=True）"""
    if force_rebuild or load_kb("labor_law") is None:
        print("正在构建法条向量库...")
        law_texts = load_law_from_docx()
        if law_texts:
            build_law_kb(law_texts, "labor_law")
            print("法条库构建完成！")
        else:
            print("⚠️ 未找到法条数据，法条库构建失败。")
    else:
        print("法条库已存在，跳过构建。")

    if force_rebuild or load_kb("labor_cases") is None:
        print("正在构建案例向量库...")
        case_texts = load_cases_from_docx()
        if case_texts:
            build_case_kb(case_texts, "labor_cases")
            print("案例库构建完成！")
        else:
            print("⚠️ 未找到案例数据，案例库构建失败。")
    else:
        print("案例库已存在，跳过构建。")

def main():
    import argparse
    parser = argparse.ArgumentParser(description="劳动法知识库管理工具")
    parser.add_argument("--init", action="store_true", help="首次初始化知识库")
    parser.add_argument("--force", action="store_true", help="强制重建知识库")
    parser.add_argument("--add-law", type=str, help="向法条库新增一条文本")
    parser.add_argument("--add-case", type=str, help="向案例库新增一条文本")
    parser.add_argument("--query", type=str, help="检索测试（法条库）")
    parser.add_argument("--query-case", type=str, help="检索测试（案例库）")
    parser.add_argument("--delete", type=str, help="按 ID 删除文档（格式：collection_name:doc_id）")

    args = parser.parse_args()

    if args.init or args.force:
        init_all_kb(force_rebuild=args.force)

    if args.add_law:
        add_docs_to_kb(args.add_law, "labor_law")
        print("法条已添加。")

    if args.add_case:
        add_docs_to_kb(args.add_case, "labor_cases")
        print("案例已添加。")

    if args.query:
        result = retrieve_from_kb(args.query, "labor_law", k=3)
        print("法条检索结果：\n", result)

    if args.query_case:
        result = retrieve_from_kb(args.query_case, "labor_cases", k=3)
        print("案例检索结果：\n", result)

    if args.delete:
        parts = args.delete.split(":")
        if len(parts) == 2:
            col, doc_id = parts
            delete_docs_from_kb([doc_id], col)
            print(f"已从 {col} 中删除文档 {doc_id}")
        else:
            print("删除格式错误，示例: labor_law:abc123")

    print("知识库构建完成！")


# ========== 命令行入口 ==========
if __name__ == "__main__":
    main()
